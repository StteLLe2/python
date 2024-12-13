import pandas as pd
from openpyxl import Workbook
from openpyxl.styles import Font, Border, Side
import json
from docx import Document

# Задание №1
def process_excel():
    files = ['file1.xlsx', 'file2.xlsx', 'file3.xlsx']
    data = pd.concat([pd.read_excel(file, header=None) for file in files])
    sorted_data = data.sort_values(by=0, ascending=False)
    sorted_data.to_excel('result.xlsx', index=False, header=False)

    wb = Workbook()
    ws = wb.active
    thin_border = Border(left=Side(style='thin'), right=Side(style='thin'),
                         top=Side(style='thin'), bottom=Side(style='thin'))
    for i, row in enumerate(sorted_data.values, 1):
        for j, cell in enumerate(row, 1):
            ws.cell(row=i, column=j, value=cell)
            ws.cell(row=i, column=j).font = Font(name='Arial', size=12)
            ws.cell(row=i, column=j).border = thin_border
    wb.save('result_styled.xlsx')

# Задание №2
def process_json():
    with open('todos.json', 'r') as f:
        todos = json.load(f)

    for todo in todos:
        with open(f"todo_{todo['id']}.json", 'w') as f:
            json.dump(todo, f, indent=4)

# Задание №3
def process_word():
    doc = Document()
    doc.add_paragraph('Hello Python', style='Normal').runs[0].bold = True
    doc.save('hello_python.docx')

    doc = Document('hello_python.docx')
    bold_text = ''.join([run.text for para in doc.paragraphs for run in para.runs if run.bold])
    print(f"Bold text: {bold_text}")

    new_doc = Document()
    para = new_doc.add_paragraph()
    run = para.add_run("This is a new paragraph with a different font and size.")
    run.font.name = 'Times New Roman'
    run.font.size = 240 
    new_doc.save('new_paragrapЁЁh.docx')

if __name__ == "__main__":
    process_excel()
    process_json()
    process_word()
